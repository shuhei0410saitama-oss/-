"""
ファイル読み込みモジュール
PDFとWord(.docx)ファイルからテキストを抽出する機能を提供
"""

from pathlib import Path


def read_pdf(file_path: str) -> str:
    """
    PDFファイルからテキストを抽出する

    Args:
        file_path: PDFファイルのパス

    Returns:
        抽出されたテキスト

    Raises:
        FileNotFoundError: ファイルが存在しない場合
        ValueError: PDFファイルでない場合
    """
    import fitz  # PyMuPDF

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"ファイルが見つかりません: {file_path}")

    if path.suffix.lower() != '.pdf':
        raise ValueError(f"PDFファイルではありません: {file_path}")

    text_parts = []

    with fitz.open(file_path) as doc:
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"--- ページ {page_num} ---\n{page_text}")

    return "\n\n".join(text_parts)


def read_docx(file_path: str) -> str:
    """
    Word(.docx)ファイルからテキストを抽出する

    Args:
        file_path: Wordファイルのパス

    Returns:
        抽出されたテキスト

    Raises:
        FileNotFoundError: ファイルが存在しない場合
        ValueError: Wordファイルでない場合
    """
    from docx import Document

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"ファイルが見つかりません: {file_path}")

    if path.suffix.lower() != '.docx':
        raise ValueError(f"Wordファイル(.docx)ではありません: {file_path}")

    doc = Document(file_path)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # テーブル内のテキストも抽出
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts)


def read_document(file_path: str) -> str:
    """
    ファイル拡張子に応じて適切な読み込み関数を呼び出す

    Args:
        file_path: ドキュメントファイルのパス

    Returns:
        抽出されたテキスト

    Raises:
        FileNotFoundError: ファイルが存在しない場合
        ValueError: サポートされていないファイル形式の場合
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"ファイルが見つかりません: {file_path}")

    extension = path.suffix.lower()

    if extension == '.pdf':
        return read_pdf(file_path)
    elif extension == '.docx':
        return read_docx(file_path)
    else:
        raise ValueError(
            f"サポートされていないファイル形式です: {extension}\n"
            "対応形式: .pdf, .docx"
        )


if __name__ == "__main__":
    # テスト用コード
    import sys

    if len(sys.argv) < 2:
        print("使用方法: python file_reader.py <ファイルパス>")
        sys.exit(1)

    file_path = sys.argv[1]

    try:
        text = read_document(file_path)
        print(f"=== {file_path} からのテキスト抽出結果 ===\n")
        print(text)
        print(f"\n=== 抽出文字数: {len(text)} ===")
    except (FileNotFoundError, ValueError) as e:
        print(f"エラー: {e}")
        sys.exit(1)
